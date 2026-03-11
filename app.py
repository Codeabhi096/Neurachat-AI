import streamlit as st
from openai import OpenAI, APITimeoutError, APIConnectionError, RateLimitError
from dotenv import load_dotenv
import datetime, re, io, os, time

try:
    from fpdf import FPDF
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

load_dotenv()

st.set_page_config(
    page_title="NeuraChat AI",
    page_icon="✦",
    layout="wide",
    initial_sidebar_state="expanded",
)

MAX_MESSAGES = 10  # Max user messages per session

# ─────────────────────────────────────────────────────────────────────────────
#  API CLIENT
# ─────────────────────────────────────────────────────────────────────────────
@st.cache_resource
def get_client():
    key = ""
    try:
        key = st.secrets["OPENROUTER_API_KEY"]
    except Exception:
        key = os.getenv("OPENROUTER_API_KEY", "")
    if not key:
        st.error(
            "**OPENROUTER_API_KEY not found!**\n\n"
            "Streamlit Cloud: App Settings → Secrets → add `OPENROUTER_API_KEY = 'sk-or-...'`\n\n"
            "Local: add to `.env` or `.streamlit/secrets.toml`"
        )
        st.stop()
    return OpenAI(base_url="https://openrouter.ai/api/v1", api_key=key, timeout=45.0)

# ─────────────────────────────────────────────────────────────────────────────
#  MODELS — Only reliable, always-available free models
# ─────────────────────────────────────────────────────────────────────────────
FREE_MODELS = [
    ("🌟 Gemini 2.0 Flash",       "google/gemini-2.0-flash-exp:free"),
    ("🧠 DeepSeek V3 0324",       "deepseek/deepseek-chat-v3-0324:free"),
    ("🦙 LLaMA 4 Maverick",       "meta-llama/llama-4-maverick:free"),
    ("🔮 Mistral Small 3.1",      "mistralai/mistral-small-3.1-24b-instruct:free"),
    ("🌙 Gemma 3 27B",            "google/gemma-3-27b-it:free"),
    ("⚡ Qwen2.5 72B",            "qwen/qwen-2.5-72b-instruct:free"),
]
FREE_MODEL_NAMES = [m[0] for m in FREE_MODELS]
FREE_MODEL_IDS   = {m[0]: m[1] for m in FREE_MODELS}

# ─────────────────────────────────────────────────────────────────────────────
#  THEMES
# ─────────────────────────────────────────────────────────────────────────────
THEMES = {
    "🌑 Midnight": {
        "bg": "#07080f", "bg2": "#0c0e1a",
        "card": "#111428", "card2": "#0f1223",
        "acc": "#6d71f0", "ahi": "#a5a8ff", "alo": "#4a4ec8",
        "soft": "rgba(109,113,240,0.10)", "glow": "rgba(109,113,240,0.22)",
        "t1": "#f0f2ff", "t2": "#8892b0", "t3": "#3d4466",
        "brd": "rgba(109,113,240,0.18)", "brd2": "rgba(109,113,240,0.32)", "brd3": "rgba(255,255,255,0.06)",
        "ubub": "linear-gradient(135deg,#3d4bda,#5865f2,#7c80f5)",
        "abub": "#111428",
        "sb": "#09091a",
        "inp": "#111428",
        "grain": "radial-gradient(ellipse 80% 60% at 15% 5%,rgba(109,113,240,0.08) 0%,transparent 60%)",
    },
    "⚡ Cyberpunk": {
        "bg": "#040408", "bg2": "#060610",
        "card": "#080818", "card2": "#06060f",
        "acc": "#00ff9f", "ahi": "#7affcb", "alo": "#00cc80",
        "soft": "rgba(0,255,159,0.08)", "glow": "rgba(0,255,159,0.20)",
        "t1": "#e8fff5", "t2": "#6affcb", "t3": "#1a4433",
        "brd": "rgba(0,255,159,0.18)", "brd2": "rgba(0,255,159,0.35)", "brd3": "rgba(0,255,159,0.08)",
        "ubub": "linear-gradient(135deg,#ff2d78,#ff6baa)",
        "abub": "#08100c",
        "sb": "#030308",
        "inp": "#080818",
        "grain": "radial-gradient(ellipse 70% 60% at 10% 10%,rgba(0,255,159,0.05) 0%,transparent 55%)",
    },
    "☀️ Nordic": {
        "bg": "#f4f6fb", "bg2": "#edf0f8",
        "card": "#ffffff", "card2": "#f8faff",
        "acc": "#4f6ef7", "ahi": "#2d4fc5", "alo": "#7b93fb",
        "soft": "rgba(79,110,247,0.08)", "glow": "rgba(79,110,247,0.18)",
        "t1": "#1a1d2e", "t2": "#4a5068", "t3": "#9aa0b8",
        "brd": "rgba(79,110,247,0.14)", "brd2": "rgba(79,110,247,0.30)", "brd3": "rgba(0,0,0,0.07)",
        "ubub": "linear-gradient(135deg,#4f6ef7,#6b84fa)",
        "abub": "#ffffff",
        "sb": "#eaecf6",
        "inp": "#ffffff",
        "grain": "radial-gradient(ellipse 80% 60% at 20% 20%,rgba(79,110,247,0.04) 0%,transparent 60%)",
    },
    "🌸 Rose": {
        "bg": "#0d080c", "bg2": "#130b10",
        "card": "#1a0d14", "card2": "#160b11",
        "acc": "#f472b6", "ahi": "#fca5d4", "alo": "#db2777",
        "soft": "rgba(244,114,182,0.10)", "glow": "rgba(244,114,182,0.22)",
        "t1": "#fff0f7", "t2": "#c77da0", "t3": "#5a2d44",
        "brd": "rgba(244,114,182,0.18)", "brd2": "rgba(244,114,182,0.32)", "brd3": "rgba(244,114,182,0.07)",
        "ubub": "linear-gradient(135deg,#db2777,#f472b6)",
        "abub": "#1a0d14",
        "sb": "#0c0810",
        "inp": "#1a0d14",
        "grain": "radial-gradient(ellipse 80% 60% at 15% 10%,rgba(244,114,182,0.07) 0%,transparent 60%)",
    },
}

# ─────────────────────────────────────────────────────────────────────────────
#  TOPIC DETECTION
# ─────────────────────────────────────────────────────────────────────────────
REF_MAP = {
    "code":    ["Stack Overflow", "GitHub", "Official Docs"],
    "math":    ["Wolfram Alpha", "ArXiv", "Khan Academy"],
    "science": ["PubMed", "Nature Journals", "arXiv"],
    "writing": ["Style Guides", "Literary Resources", "Grammarly"],
    "general": ["Wikipedia", "Web Corpus", "Academic Sources"],
    "analysis":["Research Papers", "Statistical DBs", "Industry Reports"],
    "history": ["Britannica", "Historical Archives", "Academic Journals"],
}
_CODE_KW    = {"code","python","javascript","function","bug","api","sql","html","css","def ","const ","git","react","node","docker","typescript","golang","rust","java","php","c++"}
_MATH_KW    = {"math","equation","calculus","algebra","integral","formula","statistics","matrix","derivative","proof"}
_SCI_KW     = {"science","physics","chemistry","biology","quantum","genetics","molecule","atom"}
_WRITE_KW   = {"write","essay","story","poem","email","letter","blog","creative","fiction","paragraph"}
_ANALYZE_KW = {"analyze","compare","evaluate","research","investigate","assess","explain"}
_HIST_KW    = {"history","historical","war","ancient","civilization","revolution","century"}

def detect_topic(text: str) -> str:
    t = text.lower()
    if any(w in t for w in _CODE_KW):    return "code"
    if any(w in t for w in _MATH_KW):    return "math"
    if any(w in t for w in _SCI_KW):     return "science"
    if any(w in t for w in _WRITE_KW):   return "writing"
    if any(w in t for w in _ANALYZE_KW): return "analysis"
    if any(w in t for w in _HIST_KW):    return "history"
    return "general"

def get_refs(prompt: str) -> list:
    return REF_MAP.get(detect_topic(prompt), REF_MAP["general"])[:3]

# ─────────────────────────────────────────────────────────────────────────────
#  STYLES & TONES
# ─────────────────────────────────────────────────────────────────────────────
STYLES = {
    "Balanced":  "Clear, well-structured, professional. Use markdown with headers.",
    "Concise":   "Brief and direct. Key points only. Bullet points preferred.",
    "Detailed":  "Comprehensive with examples, edge cases, and full explanations.",
    "Technical": "Precise. Always include code, formulas, implementation details.",
    "Creative":  "Vivid, imaginative, surprising. Push beyond conventional answers.",
    "Friendly":  "Warm and conversational. Explain like talking to a smart friend.",
}
TONES = ["Professional", "Friendly", "Casual", "Academic", "Creative", "Direct"]

def build_system_prompt(style: str, tone: str) -> str:
    return (
        "You are NeuraChat — a premium AI assistant for developers, researchers, and power users.\n\n"
        f"STYLE: {STYLES.get(style, STYLES['Balanced'])}\nTONE: {tone}\n\n"
        "RULES:\n"
        "- Use proper markdown: ## headers, **bold**, *italic*, `inline code`\n"
        "- Code blocks: always ```language\n"
        "- Math: $...$ inline, $$...$$ block\n"
        "- Start DIRECTLY — no preambles like 'Great question!'\n"
        "- Be accurate, concise, and genuinely helpful."
    )

# ─────────────────────────────────────────────────────────────────────────────
#  SESSION STATE
# ─────────────────────────────────────────────────────────────────────────────
_DEFAULTS = {
    "messages":      [],
    "model_key":     FREE_MODEL_NAMES[0],
    "style":         "Balanced",
    "tone":          "Professional",
    "temperature":   0.7,
    "max_tokens":    2048,
    "show_refs":     True,
    "show_tokens":   True,
    "show_timing":   True,
    "theme":         "🌑 Midnight",
    "session_start": datetime.datetime.now().strftime("%H:%M"),
    "_busy":         False,
}
for _k, _v in _DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v

# ─────────────────────────────────────────────────────────────────────────────
#  CSS — Production-ready, fully responsive, sidebar fixed
# ─────────────────────────────────────────────────────────────────────────────
def build_css(t: dict) -> str:
    return f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Söhne:wght@400;500;600&display=swap');

:root {{
  --bg:   {t["bg"]};
  --bg2:  {t["bg2"]};
  --card: {t["card"]};
  --card2:{t["card2"]};
  --acc:  {t["acc"]};
  --ahi:  {t["ahi"]};
  --alo:  {t["alo"]};
  --soft: {t["soft"]};
  --glow: {t["glow"]};
  --t1:   {t["t1"]};
  --t2:   {t["t2"]};
  --t3:   {t["t3"]};
  --brd:  {t["brd"]};
  --brd2: {t["brd2"]};
  --brd3: {t["brd3"]};
  --ubub: {t["ubub"]};
  --abub: {t["abub"]};
  --sb:   {t["sb"]};
  --inp:  {t["inp"]};
  --fd: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
  --fb: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
  --mono: 'SFMono-Regular', 'Consolas', 'Liberation Mono', 'Menlo', monospace;
  --r: 14px;
  --rs: 10px;
}}

*, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}

html, body {{ overflow-x: hidden; }}

.stApp {{
  background: var(--bg) !important;
  font-family: var(--fb) !important;
  color: var(--t1) !important;
}}

.stApp::before {{
  content: '';
  position: fixed;
  inset: 0;
  background: {t["grain"]};
  pointer-events: none;
  z-index: 0;
}}

/* Hide Streamlit chrome */
#MainMenu, footer, header, .stDeployButton,
[data-testid="stToolbar"],
[data-testid="stDecoration"],
[data-testid="stStatusWidget"] {{ display: none !important; }}

.block-container {{ padding: 0 !important; max-width: 100% !important; }}
.stAppViewBlockContainer {{ padding-top: 0 !important; }}
.stMainBlockContainer {{ padding: 0 !important; }}
.main .block-container {{ padding: 0 !important; }}

/* Scrollbar */
::-webkit-scrollbar {{ width: 3px; height: 3px; }}
::-webkit-scrollbar-track {{ background: transparent; }}
::-webkit-scrollbar-thumb {{ background: var(--acc)44; border-radius: 99px; }}
::-webkit-scrollbar-thumb:hover {{ background: var(--acc)88; }}

/* ═══════════════════════════════════════════
   SIDEBAR — FULLY FIXED (Production Safe)
   Uses solid background, no backdrop-filter,
   no transparent overrides that break layout
   ═══════════════════════════════════════════ */

section[data-testid="stSidebar"] {{
  background-color: var(--sb) !important;
  border-right: 1px solid var(--brd) !important;
  min-width: 260px !important;
  max-width: 280px !important;
  width: 270px !important;
  position: relative !important;
  z-index: 999 !important;
}}

section[data-testid="stSidebar"] > div {{
  background-color: var(--sb) !important;
  height: 100% !important;
  overflow-y: auto !important;
  overflow-x: hidden !important;
}}

[data-testid="stSidebarContent"] {{
  background-color: var(--sb) !important;
  padding: 1rem 0.85rem 3rem !important;
}}

[data-testid="stSidebarUserContent"] {{
  background-color: transparent !important;
}}

/* Sidebar collapse button */
[data-testid="stSidebarCollapseButton"] {{
  background: var(--soft) !important;
  border: 1px solid var(--brd) !important;
  border-radius: 8px !important;
  color: var(--ahi) !important;
}}
[data-testid="stSidebarCollapseButton"]:hover {{
  background: var(--glow) !important;
}}

/* Sidebar labels */
section[data-testid="stSidebar"] label,
section[data-testid="stSidebar"] .stSelectbox label,
section[data-testid="stSidebar"] .stSlider label,
section[data-testid="stSidebar"] .stToggle label {{
  color: var(--t2) !important;
  font-size: 0.8125rem !important;
  font-weight: 500 !important;
  font-family: var(--fb) !important;
}}

/* Sidebar selectbox */
section[data-testid="stSidebar"] .stSelectbox > div > div {{
  background: var(--card) !important;
  border: 1px solid var(--brd) !important;
  border-radius: var(--rs) !important;
  color: var(--t1) !important;
  font-size: 0.875rem !important;
  font-family: var(--fb) !important;
}}
section[data-testid="stSidebar"] .stSelectbox > div > div:hover {{
  border-color: var(--acc) !important;
}}
section[data-testid="stSidebar"] [data-baseweb="popover"] li {{
  background: var(--bg2) !important;
  color: var(--t1) !important;
  font-size: 0.8rem !important;
}}
section[data-testid="stSidebar"] [data-baseweb="popover"] li:hover {{
  background: var(--soft) !important;
  color: var(--ahi) !important;
}}

/* Sidebar slider */
section[data-testid="stSidebar"] [data-baseweb="thumb"] {{
  background: var(--acc) !important;
  border: 2px solid var(--bg) !important;
  box-shadow: 0 0 8px var(--glow) !important;
}}
section[data-testid="stSidebar"] [data-baseweb="track-fill"] {{
  background: linear-gradient(90deg, var(--alo), var(--acc)) !important;
}}
section[data-testid="stSidebar"] [data-baseweb="track"] {{
  background: var(--card) !important;
}}

/* Toggle */
section[data-testid="stSidebar"] [data-baseweb="checkbox"] {{
  gap: 8px !important;
}}

/* Sidebar buttons */
section[data-testid="stSidebar"] .stButton > button {{
  background: var(--soft) !important;
  border: 1px solid var(--brd2) !important;
  color: var(--ahi) !important;
  border-radius: var(--rs) !important;
  font-family: var(--fb) !important;
  font-size: 0.875rem !important;
  font-weight: 500 !important;
  width: 100% !important;
  padding: 0.5rem 0.8rem !important;
  transition: all 0.2s !important;
  cursor: pointer !important;
}}
section[data-testid="stSidebar"] .stButton > button:hover {{
  background: var(--glow) !important;
  border-color: var(--acc) !important;
  box-shadow: 0 4px 16px var(--glow) !important;
  transform: translateY(-1px) !important;
  color: var(--t1) !important;
}}

/* Download buttons */
section[data-testid="stSidebar"] [data-testid="stDownloadButton"] > button {{
  background: var(--soft) !important;
  border: 1px solid var(--brd) !important;
  color: var(--ahi) !important;
  border-radius: var(--rs) !important;
  font-family: var(--fb) !important;
  font-size: 0.73rem !important;
  font-weight: 600 !important;
  width: 100% !important;
  padding: 0.44rem 0.75rem !important;
  transition: all 0.2s !important;
}}
section[data-testid="stSidebar"] [data-testid="stDownloadButton"] > button:hover {{
  background: var(--glow) !important;
  border-color: var(--acc) !important;
  transform: translateY(-1px) !important;
}}

/* ── Sidebar Components ── */
.nc-brand {{
  display: flex;
  align-items: center;
  gap: 10px;
  padding-bottom: 1rem;
  margin-bottom: 1rem;
  border-bottom: 1px solid var(--brd3);
}}
.nc-gem {{
  width: 34px;
  height: 34px;
  flex-shrink: 0;
  background: linear-gradient(135deg, var(--alo), var(--acc), var(--ahi));
  border-radius: 10px;
  display: grid;
  place-items: center;
  font-size: 14px;
  box-shadow: 0 0 16px var(--glow);
  animation: nc-gem 4s ease-in-out infinite;
}}
@keyframes nc-gem {{
  0%,100% {{ box-shadow: 0 0 14px var(--glow); }}
  50% {{ box-shadow: 0 0 28px var(--glow); transform: scale(1.05); }}
}}
.nc-name {{
  font-family: var(--fb);
  font-size: 0.875rem;
  font-weight: 600;
  color: var(--t1);
  line-height: 1.3;
  letter-spacing: -0.01em;
}}
.nc-badge {{
  display: inline-flex;
  align-items: center;
  gap: 3px;
  background: var(--soft);
  border: 1px solid var(--brd2);
  border-radius: 99px;
  padding: 1px 7px;
  font-size: 0.6rem;
  font-weight: 500;
  color: var(--ahi);
  text-transform: none;
  letter-spacing: 0em;
  margin-top: 2px;
}}
.nc-status {{
  display: inline-flex;
  align-items: center;
  gap: 7px;
  background: var(--card);
  border: 1px solid var(--brd3);
  border-radius: 99px;
  padding: 5px 11px;
  margin-bottom: 8px;
  width: 100%;
}}
.nc-dot {{
  width: 6px;
  height: 6px;
  border-radius: 50%;
  flex-shrink: 0;
}}
.nc-on {{ background: #10b981; box-shadow: 0 0 6px #10b98199; animation: nc-blink 2.5s infinite; }}
.nc-busy {{ background: #f59e0b; box-shadow: 0 0 6px #f59e0b99; animation: nc-blink 0.6s infinite; }}
@keyframes nc-blink {{ 0%,100% {{ opacity: 1; }} 50% {{ opacity: 0.2; }} }}
.nc-stxt {{ font-size: 0.7rem; font-weight: 600; color: var(--t2); }}

.nc-lbl {{
  font-size: 0.6875rem;
  font-weight: 500;
  color: var(--t3);
  text-transform: none;
  letter-spacing: 0em;
  margin: 13px 0 5px;
  display: flex;
  align-items: center;
  gap: 4px;
}}
.nc-mchip {{
  background: var(--card);
  border: 1px solid var(--brd);
  border-radius: var(--rs);
  padding: 5px 9px;
  font-family: var(--mono);
  font-size: 0.6rem;
  color: var(--ahi);
  overflow: hidden;
  text-overflow: ellipsis;
  white-space: nowrap;
  margin-top: 3px;
}}
.nc-freebadge {{
  display: inline-block;
  padding: 1px 6px;
  border-radius: 99px;
  float: right;
  font-size: 0.625rem;
  font-weight: 500;
  background: rgba(16,185,129,0.12);
  color: #10b981;
  border: 1px solid rgba(16,185,129,0.25);
}}
.nc-stats {{
  display: grid;
  grid-template-columns: 1fr 1fr 1fr;
  gap: 4px;
  margin-top: 4px;
}}
.nc-stat {{
  background: var(--card);
  border: 1px solid var(--brd3);
  border-radius: var(--rs);
  padding: 7px 4px;
  text-align: center;
  transition: all 0.2s;
}}
.nc-stat:hover {{
  border-color: var(--acc);
  transform: translateY(-2px);
  box-shadow: 0 5px 12px var(--glow);
}}
.nc-stat-n {{
  font-family: var(--fb);
  font-size: 1.05rem;
  font-weight: 600;
  line-height: 1;
  background: linear-gradient(135deg, var(--acc), var(--ahi));
  -webkit-background-clip: text;
  -webkit-text-fill-color: transparent;
  background-clip: text;
}}
.nc-stat-l {{
  font-size: 0.5rem;
  color: var(--t3);
  margin-top: 2px;
  font-weight: 700;
  text-transform: uppercase;
  letter-spacing: 0.08em;
}}
.nc-tags {{
  display: flex;
  flex-wrap: wrap;
  gap: 4px;
  margin-top: 4px;
}}
.nc-tag {{
  background: var(--soft);
  border: 1px solid var(--brd);
  color: var(--ahi);
  border-radius: 99px;
  padding: 3px 8px;
  font-size: 0.58rem;
  font-weight: 600;
  transition: all 0.18s;
  cursor: default;
}}
.nc-tag:hover {{
  background: var(--glow);
  border-color: var(--acc);
  transform: translateY(-1px);
}}
.nc-msg-limit {{
  background: rgba(245,158,11,0.10);
  border: 1px solid rgba(245,158,11,0.30);
  border-radius: var(--r);
  padding: 8px 10px;
  margin: 8px 0;
  font-size: 0.65rem;
  color: #f59e0b;
  line-height: 1.5;
}}
.nc-footer {{
  font-size: 0.58rem;
  color: var(--t3);
  text-align: center;
  line-height: 1.8;
  margin-top: 12px;
  padding-top: 10px;
  border-top: 1px solid var(--brd3);
}}

/* ═══════════════════════════════════════════
   TOP BAR
   ═══════════════════════════════════════════ */
.nc-topbar {{
  background: var(--bg);
  border-bottom: 1px solid var(--brd3);
  padding: 0.6rem 1.4rem;
  display: flex;
  align-items: center;
  justify-content: space-between;
  position: sticky;
  top: 0;
  z-index: 100;
  flex-wrap: wrap;
  gap: 8px;
}}
.nc-tbl {{
  display: flex;
  align-items: center;
  gap: 8px;
  min-width: 0;
}}
.nc-tbico {{
  width: 26px;
  height: 26px;
  flex-shrink: 0;
  background: linear-gradient(135deg, var(--alo), var(--acc));
  border-radius: 8px;
  display: grid;
  place-items: center;
  font-size: 11px;
  box-shadow: 0 2px 8px var(--glow);
}}
.nc-tbtitle {{
  font-family: var(--fb);
  font-size: 0.88rem;
  font-weight: 600;
  color: var(--t1);
  letter-spacing: -0.01em;
  white-space: nowrap;
}}
.nc-tbr {{
  display: flex;
  align-items: center;
  gap: 5px;
  flex-wrap: wrap;
}}
.nc-pill {{
  border-radius: 99px;
  padding: 3px 9px;
  font-size: 0.6875rem;
  font-weight: 500;
  display: flex;
  align-items: center;
  gap: 3px;
  border: 1px solid var(--brd3);
  white-space: nowrap;
  transition: all 0.2s;
}}
.nc-pill:hover {{
  border-color: var(--acc);
  box-shadow: 0 2px 8px var(--glow);
}}
.nc-pm {{ background: var(--soft); color: var(--ahi); font-family: var(--mono); font-size: 0.56rem; }}
.nc-pt {{ background: var(--card); color: var(--t2); }}
.nc-ps {{ background: rgba(16,185,129,0.08); border-color: rgba(16,185,129,0.25) !important; color: #10b981; }}
.nc-pdot {{ width: 5px; height: 5px; border-radius: 50%; background: currentColor; animation: nc-blink 2s infinite; }}

/* ═══════════════════════════════════════════
   WELCOME SCREEN
   ═══════════════════════════════════════════ */
.nc-welcome {{
  display: flex;
  flex-direction: column;
  align-items: center;
  justify-content: center;
  min-height: 60vh;
  padding: 2.5rem 1rem 2rem;
  text-align: center;
  position: relative;
  z-index: 1;
}}
.nc-orb {{
  width: 78px;
  height: 78px;
  background: linear-gradient(135deg, var(--alo), var(--acc), var(--ahi));
  border-radius: 24px;
  display: grid;
  place-items: center;
  font-size: 30px;
  margin-bottom: 1.5rem;
  box-shadow: 0 0 0 1px var(--brd2), 0 8px 36px var(--glow);
  animation: nc-float 5s ease-in-out infinite;
}}
@keyframes nc-float {{
  0%,100% {{ transform: translateY(0); }}
  50% {{ transform: translateY(-9px); }}
}}
.nc-wh {{
  font-family: var(--fb);
  font-size: clamp(1.4rem, 3vw, 1.75rem);
  font-weight: 600;
  line-height: 1.15;
  color: var(--t1);
  margin-bottom: 0.6rem;
  letter-spacing: -0.025em;
}}
.nc-wh span {{
  background: linear-gradient(120deg, var(--alo), var(--acc), var(--ahi));
  -webkit-background-clip: text;
  -webkit-text-fill-color: transparent;
  background-clip: text;
}}
.nc-wsub {{
  font-size: clamp(0.85rem, 2vw, 0.9375rem);
  color: var(--t2);
  max-width: 380px;
  line-height: 1.7;
  margin-bottom: 2rem;
}}
.nc-wgrid {{
  display: grid;
  grid-template-columns: repeat(3, 1fr);
  gap: 7px;
  width: 100%;
  max-width: 540px;
}}
.nc-wcard {{
  background: var(--card);
  border: 1px solid var(--brd3);
  border-radius: var(--r);
  padding: 11px 9px;
  text-align: left;
  position: relative;
  overflow: hidden;
  transition: all 0.22s;
  cursor: default;
}}
.nc-wcard::before {{
  content: '';
  position: absolute;
  inset: 0;
  background: linear-gradient(135deg, var(--glow) 0%, transparent 100%);
  opacity: 0;
  transition: opacity 0.22s;
}}
.nc-wcard:hover {{
  border-color: var(--acc);
  transform: translateY(-3px);
  box-shadow: 0 10px 24px var(--glow);
}}
.nc-wcard:hover::before {{ opacity: 1; }}
.nc-wi {{ font-size: 1.15rem; margin-bottom: 4px; }}
.nc-wt {{ font-size: 0.75rem; font-weight: 600; color: var(--t1); margin-bottom: 2px; }}
.nc-ws {{ font-size: 0.61rem; color: var(--t2); line-height: 1.4; }}

/* ═══════════════════════════════════════════
   CHAT MESSAGES
   ═══════════════════════════════════════════ */
.nc-wrap {{
  max-width: 820px;
  margin: 0 auto;
  padding: 0.8rem clamp(0.6rem, 3vw, 1.8rem) 0.5rem;
  position: relative;
  z-index: 1;
}}

/* Hide default avatars */
[data-testid="chatAvatarIcon-user"],
[data-testid="chatAvatarIcon-assistant"] {{ display: none !important; }}

[data-testid="stChatMessage"] {{
  background: transparent !important;
  border: none !important;
  padding: 0.18rem 0 !important;
}}

/* User bubble */
[data-testid="stChatMessage"]:has([data-testid="chatAvatarIcon-user"]) .stChatMessageContent,
[data-testid="stChatMessage"][data-testid*="user"] .stChatMessageContent {{
  background: var(--ubub) !important;
  border-radius: 18px 18px 4px 18px !important;
  color: #fff !important;
  max-width: 75% !important;
  margin-left: auto !important;
  padding: 10px 14px !important;
  font-size: 0.9375rem !important;
  line-height: 1.65 !important;
  box-shadow: 0 4px 18px var(--glow) !important;
  border: none !important;
  animation: nc-msgr 0.22s ease !important;
}}

/* Assistant bubble */
[data-testid="stChatMessage"]:has([data-testid="chatAvatarIcon-assistant"]) .stChatMessageContent,
[data-testid="stChatMessage"][data-testid*="assistant"] .stChatMessageContent {{
  background: var(--abub) !important;
  border: 1px solid var(--brd3) !important;
  border-radius: 4px 18px 18px 18px !important;
  color: var(--t1) !important;
  max-width: 88% !important;
  padding: 12px 15px !important;
  font-size: 0.9375rem !important;
  line-height: 1.75 !important;
  box-shadow: 0 2px 14px rgba(0,0,0,0.18) !important;
  animation: nc-msgl 0.22s ease !important;
}}

@keyframes nc-msgr {{
  from {{ opacity: 0; transform: translateX(12px) scale(0.97); }}
  to   {{ opacity: 1; transform: none; }}
}}
@keyframes nc-msgl {{
  from {{ opacity: 0; transform: translateX(-12px) scale(0.97); }}
  to   {{ opacity: 1; transform: none; }}
}}

/* Markdown inside chat */
[data-testid="stChatMessage"] h1,
[data-testid="stChatMessage"] h2,
[data-testid="stChatMessage"] h3 {{
  font-family: var(--fd) !important;
  color: var(--ahi) !important;
  margin: 12px 0 5px !important;
  font-weight: 700 !important;
}}
[data-testid="stChatMessage"] h1 {{ font-size: 1.15em !important; border-bottom: 1px solid var(--brd3); padding-bottom: 4px !important; }}
[data-testid="stChatMessage"] h2 {{ font-size: 1.02em !important; }}
[data-testid="stChatMessage"] h3 {{ font-size: 0.93em !important; }}
[data-testid="stChatMessage"] p  {{ margin-bottom: 5px !important; }}
[data-testid="stChatMessage"] ul,
[data-testid="stChatMessage"] ol {{ padding-left: 16px !important; margin: 4px 0 !important; }}
[data-testid="stChatMessage"] li {{ margin-bottom: 3px !important; color: var(--t2) !important; }}
[data-testid="stChatMessage"] strong {{ color: var(--t1) !important; font-weight: 700 !important; }}
[data-testid="stChatMessage"] em {{ color: var(--ahi) !important; }}
[data-testid="stChatMessage"] a  {{ color: var(--ahi) !important; text-decoration: underline !important; }}

/* Code */
[data-testid="stChatMessage"] code {{
  background: var(--soft) !important;
  border: 1px solid var(--brd) !important;
  border-radius: 5px !important;
  padding: 2px 6px !important;
  font-size: 0.81em !important;
  color: var(--ahi) !important;
  font-family: var(--mono) !important;
}}
[data-testid="stChatMessage"] pre {{
  background: rgba(5,7,14,0.95) !important;
  border: 1px solid var(--brd) !important;
  border-left: 3px solid var(--acc) !important;
  border-radius: 10px !important;
  padding: 12px !important;
  overflow-x: auto !important;
  margin: 9px 0 !important;
}}
[data-testid="stChatMessage"] pre code {{
  background: transparent !important;
  border: none !important;
  padding: 0 !important;
  color: var(--t1) !important;
  font-size: 0.82em !important;
}}

/* Tables */
[data-testid="stChatMessage"] table {{
  border-collapse: collapse !important;
  width: 100% !important;
  margin: 9px 0 !important;
  font-size: 0.84rem !important;
}}
[data-testid="stChatMessage"] th {{
  background: var(--soft) !important;
  color: var(--ahi) !important;
  padding: 6px 10px !important;
  font-size: 0.69rem !important;
  font-weight: 700 !important;
  text-transform: uppercase !important;
  border-bottom: 1px solid var(--brd) !important;
}}
[data-testid="stChatMessage"] td {{
  padding: 6px 10px !important;
  border-bottom: 1px solid var(--brd3) !important;
  color: var(--t2) !important;
}}
[data-testid="stChatMessage"] tr:hover td {{
  background: var(--soft) !important;
  color: var(--t1) !important;
}}
[data-testid="stChatMessage"] blockquote {{
  border-left: 3px solid var(--acc) !important;
  margin: 8px 0 !important;
  padding: 7px 12px !important;
  background: var(--soft) !important;
  border-radius: 0 9px 9px 0 !important;
  color: var(--t2) !important;
  font-style: italic !important;
}}

/* Meta strip */
.nc-meta {{
  display: flex;
  flex-wrap: wrap;
  align-items: center;
  gap: 4px;
  margin-top: 6px;
}}
.nc-chip {{
  display: inline-flex;
  align-items: center;
  gap: 3px;
  background: var(--card);
  border: 1px solid var(--brd3);
  border-radius: 99px;
  padding: 2px 8px;
  font-size: 0.6875rem;
  font-weight: 400;
  color: var(--t3);
  white-space: nowrap;
}}
.nc-chip span {{ color: var(--t2); }}
.nc-refs {{
  display: flex;
  flex-wrap: wrap;
  align-items: center;
  gap: 4px;
  margin-top: 5px;
  padding: 5px 10px;
  background: var(--soft);
  border: 1px solid var(--brd3);
  border-radius: var(--rs);
}}
.nc-refs-lbl {{
  font-size: 0.56rem;
  font-weight: 700;
  color: var(--t3);
  text-transform: uppercase;
  letter-spacing: 0.1em;
  margin-right: 2px;
}}
.nc-ref {{
  background: var(--card);
  border: 1px solid var(--brd);
  border-radius: 99px;
  padding: 2px 8px;
  font-size: 0.59rem;
  font-weight: 600;
  color: var(--ahi);
  white-space: nowrap;
  transition: all 0.18s;
}}
.nc-ref:hover {{
  background: var(--glow);
  transform: translateY(-1px);
}}

/* Typing / generating */
.nc-typing {{
  display: flex;
  align-items: center;
  gap: 5px;
  padding: 9px 13px;
  background: var(--abub);
  border: 1px solid var(--brd3);
  border-radius: 4px 16px 16px 16px;
  width: fit-content;
  margin: 3px 0;
  box-shadow: 0 2px 10px rgba(0,0,0,0.18);
}}
.nc-td {{
  width: 5px;
  height: 5px;
  border-radius: 50%;
  background: var(--acc);
  animation: nc-tdot 1.3s ease-in-out infinite;
}}
.nc-td:nth-child(1) {{ animation-delay: 0s; }}
.nc-td:nth-child(2) {{ animation-delay: 0.22s; }}
.nc-td:nth-child(3) {{ animation-delay: 0.44s; }}
@keyframes nc-tdot {{
  0%,55%,100% {{ opacity: 0.15; transform: translateY(0); }}
  28% {{ opacity: 1; transform: translateY(-4px); }}
}}
.nc-tlbl {{
  font-size: 0.65rem;
  color: var(--t3);
  font-style: italic;
  margin-left: 3px;
}}
.nc-gen {{
  display: inline-flex;
  align-items: center;
  gap: 6px;
  background: var(--soft);
  border: 1px solid var(--brd2);
  border-radius: 99px;
  padding: 4px 11px;
  font-size: 0.67rem;
  font-weight: 600;
  color: var(--ahi);
  margin-bottom: 7px;
  animation: nc-pulse 1.2s ease-in-out infinite;
}}
@keyframes nc-pulse {{
  0%,100% {{ box-shadow: 0 0 0 0 var(--glow); }}
  50% {{ box-shadow: 0 0 0 4px transparent; }}
}}
.nc-gd {{
  width: 6px;
  height: 6px;
  border-radius: 50%;
  background: var(--acc);
  animation: nc-gda 0.9s ease-in-out infinite alternate;
}}
@keyframes nc-gda {{
  from {{ opacity: 0.3; transform: scale(0.7); }}
  to   {{ opacity: 1; transform: scale(1.2); }}
}}

/* Limit warning banner */
.nc-limit-banner {{
  max-width: 820px;
  margin: 1rem auto;
  padding: 1.2rem 1.5rem;
  background: rgba(245,158,11,0.08);
  border: 1px solid rgba(245,158,11,0.30);
  border-radius: var(--r);
  text-align: center;
  position: relative;
  z-index: 1;
}}
.nc-limit-banner h3 {{
  font-family: var(--fb);
  font-size: 1rem;
  color: #f59e0b;
  margin-bottom: 5px;
  font-weight: 700;
}}
.nc-limit-banner p {{
  font-size: 0.82rem;
  color: var(--t2);
  line-height: 1.6;
}}

/* Error banner */
.nc-error-banner {{
  background: rgba(239,68,68,0.08);
  border: 1px solid rgba(239,68,68,0.25);
  border-radius: var(--r);
  padding: 10px 14px;
  margin: 6px 0;
  font-size: 0.8rem;
  color: #fca5a5;
  line-height: 1.55;
}}

/* ═══════════════════════════════════════════
   CHAT INPUT
   ═══════════════════════════════════════════ */
[data-testid="stBottom"] {{
  background: transparent !important;
  border-top: none !important;
  padding: 0.6rem clamp(0.6rem, 3.5vw, 1.6rem) 0.8rem !important;
  position: sticky !important;
  bottom: 0 !important;
  z-index: 100 !important;
}}
[data-testid="stBottom"]::before {{
  content: '';
  position: absolute;
  inset: 0;
  background: linear-gradient(to top, var(--bg) 50%, transparent);
  border-top: 1px solid var(--brd3);
  z-index: -1;
}}
[data-testid="stChatInput"] {{
  background: var(--inp) !important;
  border: 1.5px solid var(--brd2) !important;
  border-radius: 18px !important;
  max-width: 800px !important;
  margin: 0 auto !important;
  transition: border-color 0.25s, box-shadow 0.25s !important;
  box-shadow: 0 4px 18px rgba(0,0,0,0.20) !important;
}}
[data-testid="stChatInput"]:focus-within {{
  border-color: var(--acc) !important;
  box-shadow: 0 0 0 3px var(--glow), 0 4px 20px var(--glow) !important;
}}
[data-testid="stChatInput"] textarea {{
  background: transparent !important;
  color: var(--t1) !important;
  font-family: var(--fb) !important;
  font-size: 0.9375rem !important;
  caret-color: var(--ahi) !important;
  padding: 12px 15px !important;
  line-height: 1.55 !important;
  min-height: 48px !important;
}}
[data-testid="stChatInput"] textarea::placeholder {{
  color: var(--t3) !important;
  font-size: 0.9375rem !important;
}}
[data-testid="stChatInput"] button {{
  background: linear-gradient(135deg, var(--alo), var(--acc)) !important;
  border: none !important;
  border-radius: 11px !important;
  margin: 5px !important;
  transition: opacity 0.18s, transform 0.18s !important;
  box-shadow: 0 2px 10px var(--glow) !important;
}}
[data-testid="stChatInput"] button:hover {{
  opacity: 0.85 !important;
  transform: scale(1.08) !important;
}}
[data-testid="stChatInput"] button svg {{ fill: #fff !important; }}

/* Disabled input state */
[data-testid="stChatInput"][disabled],
[data-testid="stChatInput"].disabled {{
  opacity: 0.5 !important;
  pointer-events: none !important;
}}

/* Misc */
hr {{ border: none !important; border-top: 1px solid var(--brd3) !important; margin: 9px 0 !important; }}

/* ═══════════════════════════════════════════
   RESPONSIVE — Mobile First
   ═══════════════════════════════════════════ */
@media (max-width: 768px) {{
  .nc-topbar {{
    padding: 0.5rem 0.8rem;
    gap: 6px;
  }}
  .nc-pm {{ display: none !important; }}
  .nc-tbr .nc-pill:nth-child(3) {{ display: none !important; }}
  section[data-testid="stSidebar"] {{
    min-width: 240px !important;
    max-width: 260px !important;
    width: 250px !important;
  }}
  .nc-wh {{ font-size: clamp(1.3rem, 6vw, 1.6rem); }}
  .nc-wgrid {{ grid-template-columns: 1fr 1fr; gap: 6px; }}
  .nc-orb {{ width: 62px; height: 62px; font-size: 24px; }}
  .nc-stats {{ grid-template-columns: 1fr 1fr; }}
  [data-testid="stBottom"] {{ padding: 0.45rem 0.6rem 0.7rem !important; }}
  [data-testid="stChatInput"] {{ border-radius: 14px !important; }}
  [data-testid="stChatMessage"]:has([data-testid="chatAvatarIcon-user"]) .stChatMessageContent {{
    max-width: 88% !important;
  }}
  [data-testid="stChatMessage"]:has([data-testid="chatAvatarIcon-assistant"]) .stChatMessageContent {{
    max-width: 95% !important;
  }}
  .nc-wrap {{ padding: 0.6rem 0.7rem 0.4rem; }}
}}

@media (max-width: 480px) {{
  .nc-wgrid {{ grid-template-columns: 1fr; max-width: 280px; }}
  .nc-pill {{ padding: 3px 7px; font-size: 0.57rem; }}
  .nc-tbr .nc-pill:nth-child(2) {{ display: none !important; }}
  .nc-tbtitle {{ font-size: 0.82rem; }}
  section[data-testid="stSidebar"] {{
    min-width: 220px !important;
    max-width: 240px !important;
    width: 230px !important;
  }}
}}

/* Spinner */
.stSpinner > div {{ border-top-color: var(--acc) !important; }}
</style>
"""

# ─────────────────────────────────────────────────────────────────────────────
#  EXPORT HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def export_txt(messages: list) -> bytes:
    lines = [
        "NeuraChat AI — Conversation Export",
        f"Date  : {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}",
        f"Model : {FREE_MODEL_IDS.get(st.session_state.model_key, 'unknown')}",
        "═" * 60, "",
    ]
    for m in messages:
        lines += [f"[{'You' if m['role'] == 'user' else 'NeuraChat AI'}]", m["content"], ""]
    return "\n".join(lines).encode("utf-8")

def export_md(messages: list) -> bytes:
    model = FREE_MODEL_IDS.get(st.session_state.model_key, "unknown")
    lines = ["# NeuraChat AI — Conversation Export",
             f"*{datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}*  ·  Model: `{model}`", ""]
    for m in messages:
        role = "**You**" if m["role"] == "user" else "**NeuraChat AI**"
        lines += [f"### {role}", m["content"], "---", ""]
    return "\n".join(lines).encode("utf-8")

def export_pdf(messages: list) -> bytes:
    class PDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "B", 16)
            self.set_text_color(109, 113, 240)
            self.cell(0, 10, "NeuraChat AI", ln=False, align="C")
            self.ln(7)
            self.set_font("Helvetica", "", 8)
            self.set_text_color(140, 145, 170)
            model = FREE_MODEL_IDS.get(st.session_state.model_key, "?")
            self.cell(0, 5, f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}  ·  {model}", ln=True, align="C")
            self.ln(3)
            self.set_draw_color(109, 113, 240)
            self.set_line_width(0.4)
            self.line(10, self.get_y(), self.w - 10, self.get_y())
            self.ln(5)
        def footer(self):
            self.set_y(-12)
            self.set_font("Helvetica", "I", 7)
            self.set_text_color(150, 150, 170)
            self.cell(0, 8, f"Page {self.page_no()} — NeuraChat AI", align="C")

    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    for m in messages:
        is_user = m["role"] == "user"
        pdf.set_font("Helvetica", "B", 10)
        if is_user:
            pdf.set_text_color(60, 80, 220)
            pdf.set_fill_color(240, 242, 255)
        else:
            pdf.set_text_color(109, 113, 240)
            pdf.set_fill_color(245, 245, 255)
        pdf.set_draw_color(200, 205, 250)
        pdf.set_line_width(0.2)
        pdf.rect(10, pdf.get_y(), pdf.w - 20, 8, "DF")
        pdf.set_xy(10, pdf.get_y() + 1.5)
        pdf.cell(pdf.w - 20, 5, "  YOU" if is_user else "  NEURACHAT AI", ln=True)
        pdf.ln(2)
        pdf.set_font("Helvetica", "", 9.5)
        pdf.set_text_color(30, 34, 60)
        clean = re.sub(r"```[\w]*\n?", "", m["content"])
        clean = re.sub(r"[`*#_\[\]>]+", "", clean)
        clean = re.sub(r"\n{3,}", "\n\n", clean.strip())
        safe = "".join(c if c.encode("latin-1", errors="ignore") else "?" for c in clean)
        pdf.multi_cell(0, 5.6, safe, border=0)
        pdf.ln(4)
        pdf.set_draw_color(220, 222, 235)
        pdf.line(10, pdf.get_y(), pdf.w - 10, pdf.get_y())
        pdf.ln(5)
    return bytes(pdf.output())

def export_docx(messages: list) -> bytes:
    doc = DocxDocument()
    h = doc.add_heading("NeuraChat AI — Conversation Export", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(109, 113, 240)
    sub = doc.add_paragraph(f"Exported: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}  ·  Model: {FREE_MODEL_IDS.get(st.session_state.model_key, '?')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        sub.runs[0].font.size = Pt(9)
        sub.runs[0].font.color.rgb = RGBColor(130, 130, 150)
    doc.add_paragraph()
    for m in messages:
        is_user = m["role"] == "user"
        p = doc.add_paragraph()
        rr = p.add_run(f"[{'You' if is_user else 'NeuraChat AI'}]")
        rr.bold = True
        rr.font.size = Pt(10)
        rr.font.color.rgb = RGBColor(50, 50, 80) if is_user else RGBColor(109, 113, 240)
        clean = re.sub(r"[`*#_]+", "", m["content"])
        dp = doc.add_paragraph(clean)
        if dp.runs:
            dp.runs[0].font.size = Pt(10)
        doc.add_paragraph()
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ─────────────────────────────────────────────────────────────────────────────
#  STREAMING — Smart fallback with friendly error messages
# ─────────────────────────────────────────────────────────────────────────────
def stream_response(messages: list, model_key: str, temperature: float, max_tokens: int):
    client  = get_client()
    primary = FREE_MODEL_IDS.get(model_key, FREE_MODELS[0][1])
    all_ids = [m[1] for m in FREE_MODELS]
    # Always try primary first, then fallback chain
    cands   = [primary] + [mid for mid in all_ids if mid != primary]

    api_msgs = [
        {"role": "system", "content": build_system_prompt(st.session_state.style, st.session_state.tone)}
    ] + [{"role": m["role"], "content": m["content"]} for m in messages]

    last_error = "Unknown error"
    tried = 0
    for model in cands:
        tried += 1
        try:
            stream = client.chat.completions.create(
                model=model,
                messages=api_msgs,
                max_tokens=max_tokens,
                temperature=temperature,
                stream=True,
                extra_headers={
                    "HTTP-Referer": "https://neurachat.app",
                    "X-Title": "NeuraChat AI",
                },
            )
            yielded = False
            for chunk in stream:
                d = chunk.choices[0].delta if chunk.choices else None
                if d and d.content:
                    yield d.content
                    yielded = True
            if yielded:
                return
            # Empty response — try next
            continue

        except APITimeoutError:
            last_error = "timeout"
            if model != cands[-1]:
                continue
        except RateLimitError as e:
            last_error = str(e)
            if model != cands[-1]:
                continue
        except APIConnectionError as e:
            yield (
                "\n\n**⚠️ Network Error**\n\n"
                "Internet connection issue. Please check your connection and try again."
            )
            return
        except Exception as e:
            last_error = str(e)
            err = last_error.lower()
            if any(k in err for k in ["429", "404", "quota", "not found", "temporarily",
                                       "overloaded", "unavailable", "no endpoints",
                                       "moderation", "context length"]):
                if model != cands[-1]:
                    continue
            else:
                yield (
                    f"\n\n**⚠️ Unexpected Error**\n\n"
                    f"`{last_error[:200]}`\n\nPlease try again in a moment."
                )
                return

    # All models failed
    yield (
        "\n\n**🕐 Servers are busy right now**\n\n"
        "All AI models are currently at capacity. This usually resolves within **1–2 minutes**.\n\n"
        "Please wait a moment and try your message again. "
        "If the issue persists, your API credits may be exhausted — "
        "check your [OpenRouter dashboard](https://openrouter.ai/account)."
    )

# ─────────────────────────────────────────────────────────────────────────────
#  INJECT CSS
# ─────────────────────────────────────────────────────────────────────────────
_th = THEMES[st.session_state.theme]
st.markdown(build_css(_th), unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
#  SIDEBAR
# ─────────────────────────────────────────────────────────────────────────────
with st.sidebar:
    _busy = st.session_state.get("_busy", False)
    _user_msgs = sum(1 for m in st.session_state.messages if m["role"] == "user")
    _msgs_left = MAX_MESSAGES - _user_msgs

    st.markdown(f"""
<div class="nc-brand">
  <div class="nc-gem">✦</div>
  <div>
    <div class="nc-name">NeuraChat AI</div>
    <div class="nc-badge">✦ Free · Unlimited</div>
  </div>
</div>
<div class="nc-status">
  <div class="nc-dot {'nc-busy' if _busy else 'nc-on'}"></div>
  <span class="nc-stxt">{'Generating response…' if _busy else 'Ready · All Models Active'}</span>
</div>
""", unsafe_allow_html=True)

    if _msgs_left <= 3 and _msgs_left > 0:
        st.markdown(f'<div class="nc-msg-limit">⚠️ <b>{_msgs_left} message{"s" if _msgs_left > 1 else ""} remaining</b> in this session.<br>Clear chat to reset.</div>', unsafe_allow_html=True)
    elif _msgs_left <= 0:
        st.markdown('<div class="nc-msg-limit">🔒 <b>Session limit reached</b><br>Clear the conversation to start a new session.</div>', unsafe_allow_html=True)

    # Theme
    st.markdown('<div class="nc-lbl">🎨 Theme</div>', unsafe_allow_html=True)
    _tlist = list(THEMES.keys())
    _new_theme = st.selectbox("Theme", _tlist,
                              index=_tlist.index(st.session_state.theme),
                              label_visibility="collapsed", key="sb_theme")
    if _new_theme != st.session_state.theme:
        st.session_state.theme = _new_theme
        st.rerun()

    # Model
    st.markdown('<div class="nc-lbl">🤖 AI Model</div>', unsafe_allow_html=True)
    _mi = FREE_MODEL_NAMES.index(st.session_state.model_key) \
          if st.session_state.model_key in FREE_MODEL_NAMES else 0
    st.session_state.model_key = st.selectbox(
        "Model", FREE_MODEL_NAMES, index=_mi,
        label_visibility="collapsed", key="sb_model")
    _sid = FREE_MODEL_IDS.get(st.session_state.model_key, "").split("/")[-1].replace(":free", "")
    st.markdown(f"""
<div class="nc-mchip">⚡ {_sid}<span class="nc-freebadge">FREE</span></div>
<div style="font-size:0.58rem;color:var(--t3);margin-top:3px;">Auto-fallback to next model on failure</div>
""", unsafe_allow_html=True)

    # Generation
    st.markdown('<div class="nc-lbl">⚙️ Generation</div>', unsafe_allow_html=True)
    st.session_state.temperature = st.slider("Temperature", 0.0, 1.0,
                                              float(st.session_state.temperature), 0.05,
                                              key="sb_temp", help="Higher = more creative")
    st.session_state.max_tokens  = st.slider("Max Tokens", 256, 4096,
                                              int(st.session_state.max_tokens), 64,
                                              key="sb_tok", help="Max response length")

    # Style & Tone
    st.markdown('<div class="nc-lbl">📝 Style & Tone</div>', unsafe_allow_html=True)
    st.session_state.style = st.selectbox("Style", list(STYLES.keys()),
        index=list(STYLES.keys()).index(st.session_state.style),
        key="sb_style")
    st.session_state.tone  = st.selectbox("Tone", TONES,
        index=TONES.index(st.session_state.tone) if st.session_state.tone in TONES else 0,
        key="sb_tone")

    # Options
    st.markdown('<div class="nc-lbl">🔧 Display Options</div>', unsafe_allow_html=True)
    st.session_state.show_refs   = st.toggle("📎 Source References", value=st.session_state.show_refs,   key="sb_refs")
    st.session_state.show_tokens = st.toggle("📊 Token Estimate",    value=st.session_state.show_tokens, key="sb_tkest")
    st.session_state.show_timing = st.toggle("⏱️ Response Time",     value=st.session_state.show_timing, key="sb_time")

    # Stats
    st.markdown('<div class="nc-lbl">📊 Session Stats</div>', unsafe_allow_html=True)
    _msgs    = st.session_state.messages
    _uc      = sum(1 for m in _msgs if m["role"] == "user")
    _ac      = sum(1 for m in _msgs if m["role"] == "assistant")
    _tw      = sum(len(m["content"].split()) for m in _msgs)
    _timings = [m["timing"] for m in _msgs if m.get("timing")]
    _avgt    = sum(_timings) / len(_timings) if _timings else 0
    st.markdown(f"""
<div class="nc-stats">
  <div class="nc-stat"><div class="nc-stat-n">{_uc}</div><div class="nc-stat-l">Sent</div></div>
  <div class="nc-stat"><div class="nc-stat-n">{_ac}</div><div class="nc-stat-l">Replies</div></div>
  <div class="nc-stat"><div class="nc-stat-n">{_tw}</div><div class="nc-stat-l">Words</div></div>
</div>""", unsafe_allow_html=True)
    if _avgt:
        st.markdown(f'<div style="font-size:0.6rem;color:var(--t3);margin-top:4px;">Avg response: <span style="color:var(--t2)">{_avgt:.1f}s</span></div>', unsafe_allow_html=True)

    # Capabilities
    st.markdown('<div class="nc-lbl">✨ Capabilities</div>', unsafe_allow_html=True)
    st.markdown("""<div class="nc-tags">
  <span class="nc-tag">💻 Code</span>
  <span class="nc-tag">📊 Data</span>
  <span class="nc-tag">🧮 Math</span>
  <span class="nc-tag">✍️ Writing</span>
  <span class="nc-tag">🔬 Science</span>
  <span class="nc-tag">🎨 Creative</span>
  <span class="nc-tag">📈 Analysis</span>
  <span class="nc-tag">🗺️ Diagrams</span>
</div>""", unsafe_allow_html=True)

    # Export
    st.markdown('<div class="nc-lbl">💾 Export Chat</div>', unsafe_allow_html=True)
    if _msgs:
        _fn = f"neurachat_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}"
        st.download_button("📄 Export as Text",     data=export_txt(_msgs), file_name=f"{_fn}.txt",  mime="text/plain",    key="dl_txt")
        st.download_button("📝 Export as Markdown", data=export_md(_msgs),  file_name=f"{_fn}.md",   mime="text/markdown", key="dl_md")
        if HAS_PDF:
            try:
                st.download_button("📕 Export as PDF",  data=export_pdf(_msgs),  file_name=f"{_fn}.pdf",  mime="application/pdf", key="dl_pdf")
            except Exception as e:
                st.caption(f"⚠️ PDF unavailable: {e}")
        if HAS_DOCX:
            try:
                st.download_button("📘 Export as Word", data=export_docx(_msgs), file_name=f"{_fn}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_docx")
            except Exception as e:
                st.caption(f"⚠️ DOCX unavailable: {e}")
    else:
        st.markdown('<span style="font-size:.7rem;color:var(--t3)">Start chatting to enable export</span>', unsafe_allow_html=True)

    st.markdown("---")
    if st.button("🗑️ Clear Conversation", key="btn_clear"):
        st.session_state.messages = []
        st.session_state._busy = False
        st.rerun()

    st.markdown(f"""<div class="nc-footer">
  ✦ NeuraChat AI · Free Unlimited<br>
  Powered by OpenRouter · 6 Models<br>
  Session started {st.session_state.session_start}
</div>""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
#  MAIN AREA
# ─────────────────────────────────────────────────────────────────────────────
_ms   = FREE_MODEL_IDS.get(st.session_state.model_key, "").split("/")[-1].replace(":free", "")
_tlbl = st.session_state.theme.split()[0]

st.markdown(f"""
<div class="nc-topbar">
  <div class="nc-tbl">
    <div class="nc-tbico">✦</div>
    <div class="nc-tbtitle">NeuraChat AI</div>
  </div>
  <div class="nc-tbr">
    <div class="nc-pill nc-pm">⚡ {_ms}</div>
    <div class="nc-pill nc-pt">📝 {st.session_state.style}</div>
    <div class="nc-pill nc-pt">{_tlbl}</div>
    <div class="nc-pill nc-ps"><div class="nc-pdot"></div>Free Unlimited</div>
  </div>
</div>
""", unsafe_allow_html=True)

# Check message limit
_user_count = sum(1 for m in st.session_state.messages if m["role"] == "user")
_limit_hit  = _user_count >= MAX_MESSAGES

# Welcome screen
if not st.session_state.messages:
    st.markdown("""
<div class="nc-welcome">
  <div class="nc-orb">✦</div>
  <div class="nc-wh">Hello! What shall we<br><span>explore today?</span></div>
  <div class="nc-wsub">Unlimited free AI — 6 models, auto-fallback, no limits.<br>Code, math, writing, research, and beyond.</div>
  <div class="nc-wgrid">
    <div class="nc-wcard"><div class="nc-wi">💻</div><div class="nc-wt">Code & Debug</div><div class="nc-ws">Any language, architecture, bug fixes</div></div>
    <div class="nc-wcard"><div class="nc-wi">📊</div><div class="nc-wt">Diagrams</div><div class="nc-ws">Mermaid, flowcharts, ERDs</div></div>
    <div class="nc-wcard"><div class="nc-wi">🧮</div><div class="nc-wt">Math & LaTeX</div><div class="nc-ws">Equations, proofs, step-by-step</div></div>
    <div class="nc-wcard"><div class="nc-wi">✍️</div><div class="nc-wt">Writing</div><div class="nc-ws">Reports, emails, essays, blogs</div></div>
    <div class="nc-wcard"><div class="nc-wi">🔍</div><div class="nc-wt">Research</div><div class="nc-ws">Deep analysis, summaries, compare</div></div>
    <div class="nc-wcard"><div class="nc-wi">🎨</div><div class="nc-wt">Creative</div><div class="nc-ws">Brainstorm, fiction, worldbuilding</div></div>
  </div>
</div>""", unsafe_allow_html=True)

# Chat history
with st.container():
    st.markdown('<div class="nc-wrap">', unsafe_allow_html=True)
    for _msg in st.session_state.messages:
        with st.chat_message(_msg["role"]):
            st.markdown(_msg["content"])
            if _msg["role"] == "assistant":
                _wc   = len(_msg["content"].split())
                _meta = [f'<div class="nc-chip">📝 <span>{_wc} words</span></div>']
                if st.session_state.show_tokens:
                    _meta.append(f'<div class="nc-chip">🔢 <span>~{int(_wc * 1.35)} tokens</span></div>')
                if st.session_state.show_timing and _msg.get("timing"):
                    _meta.append(f'<div class="nc-chip">⏱️ <span>{_msg["timing"]:.1f}s</span></div>')
                st.markdown(f'<div class="nc-meta">{"".join(_meta)}</div>', unsafe_allow_html=True)
                if st.session_state.show_refs and _msg.get("refs"):
                    _pills = "".join(f'<span class="nc-ref">📎 {r}</span>' for r in _msg["refs"])
                    st.markdown(f'<div class="nc-refs"><span class="nc-refs-lbl">Sources</span>{_pills}</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

# Session limit banner
if _limit_hit:
    st.markdown(f"""
<div class="nc-limit-banner">
  <h3>🔒 Session Limit Reached ({MAX_MESSAGES} messages)</h3>
  <p>
    You've used all <b>{MAX_MESSAGES} messages</b> in this session.<br>
    Click <b>"🗑️ Clear Conversation"</b> in the sidebar to start a fresh session.
  </p>
</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
#  INPUT + STREAMING
# ─────────────────────────────────────────────────────────────────────────────
if not _limit_hit:
    _placeholder = f"Ask NeuraChat anything… ({st.session_state.style} · {_ms})"
    if _prompt := st.chat_input(_placeholder):
        _refs = get_refs(_prompt) if st.session_state.show_refs else []
        st.session_state._busy = True
        st.session_state.messages.append({"role": "user", "content": _prompt})

        with st.chat_message("user"):
            st.markdown(_prompt)

        with st.chat_message("assistant"):
            _gph = st.empty()
            _tph = st.empty()
            _rph = st.empty()

            _gph.markdown(
                '<div class="nc-gen"><div class="nc-gd"></div>Generating…</div>',
                unsafe_allow_html=True
            )
            _tph.markdown(
                '<div class="nc-typing"><div class="nc-td"></div><div class="nc-td"></div><div class="nc-td"></div><span class="nc-tlbl">Thinking…</span></div>',
                unsafe_allow_html=True
            )

            _reply = ""
            _first = True
            _buf   = 0
            _t0    = time.time()

            for _chunk in stream_response(
                st.session_state.messages,
                st.session_state.model_key,
                st.session_state.temperature,
                st.session_state.max_tokens,
            ):
                if _first:
                    _gph.empty()
                    _tph.empty()
                    _first = False
                _reply += _chunk
                _buf   += 1
                if _buf >= 5:
                    _rph.markdown(_reply + "▌")
                    _buf = 0

            _elapsed = time.time() - _t0

            if _first:  # Nothing was yielded at all
                _gph.empty()
                _tph.empty()

            _rph.markdown(_reply)

            _wc   = len(_reply.split())
            _meta = [f'<div class="nc-chip">📝 <span>{_wc} words</span></div>']
            if st.session_state.show_tokens:
                _meta.append(f'<div class="nc-chip">🔢 <span>~{int(_wc * 1.35)} tokens</span></div>')
            if st.session_state.show_timing:
                _meta.append(f'<div class="nc-chip">⏱️ <span>{_elapsed:.1f}s</span></div>')
            st.markdown(f'<div class="nc-meta">{"".join(_meta)}</div>', unsafe_allow_html=True)

            if _refs:
                _pills = "".join(f'<span class="nc-ref">📎 {r}</span>' for r in _refs)
                st.markdown(
                    f'<div class="nc-refs"><span class="nc-refs-lbl">Sources</span>{_pills}</div>',
                    unsafe_allow_html=True
                )

        st.session_state._busy = False
        st.session_state.messages.append({
            "role": "assistant",
            "content": _reply,
            "refs": _refs,
            "timing": _elapsed,
        })
        st.rerun()